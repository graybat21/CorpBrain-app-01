"""텍스트 추출 — 지원 포맷의 확장자 디스패치 (스펙 §4.2·§4.1).

- `.txt` / `.md`: 원문을 특별 파싱 없이 평문으로 읽는다 (FR-006).
- `.docx`: 문단 텍스트를 순서대로 이어 붙인다 (FR-007).
- `.pdf`: 텍스트 레이어를 페이지 순서로 이어 붙인다 — 이미지/OCR·암호화는 다루지 않는다
  (v0.2 U1, 스펙 §4.1).

공통 규칙: 파일을 통째로 메모리에 올리지 않고 앞부분 `max_chars`까지만 읽는다.
개별 파일의 추출 실패는 `ExtractionError`로 올리고, 호출자(FR-008/FR-015)가 스킵으로
흡수해 나머지 파일 처리를 계속한다 — 전체 실패로 위장하지 않는다 (스펙 §5).
`PermissionError`로 인한 실패는 `ExtractionError.__cause__`에 원인이 보존되므로,
호출자가 `permission_denied`와 `extraction_failed`를 구분할 수 있다.
"""

from __future__ import annotations

from collections.abc import Callable, Iterable
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from pypdf import PdfReader
from pypdf.errors import PyPdfError

from corpbrain.core.errors import CorpBrainError
from corpbrain.core.models import SkippedFile, SkipReason

#: OLE 복합문서(CFBF)의 매직 넘버 8바이트 (스펙 §4.3.1).
#:
#: 암호로 보호된 `.xlsx`/`.pptx`는 OOXML zip이 **아니라** 이 컨테이너로 감싸져 저장되므로,
#: `openpyxl`·`python-pptx`는 「zip을 열지 못했다」는 예외만 던진다 — 손상된 파일과 예외 종류가
#: 같아 그것만으로는 원인을 가를 수 없다. 구형 `.xls`/`.ppt`(BIFF·OLE)도 같은 시그니처를 갖는다.
OLE_SIGNATURE: bytes = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


class ExtractionError(CorpBrainError):
    """텍스트 추출 실패 — 해당 파일만 스킵되고 나머지 처리는 계속된다."""


@dataclass(frozen=True)
class PreparedDocument:
    """요약 입력 준비 결과 — `text`(정상)와 `skipped`(스킵) 중 정확히 하나만 채워진다."""

    text: str | None = None
    skipped: SkippedFile | None = None


def prepare_summary_input(path: Path, max_chars: int) -> PreparedDocument:
    """추출 → 길이 제한 → 빈 문서·실패 판정을 한 규칙으로 묶는다 (FR-008 / 스펙 §4.2·§5).

    12,000자(기본) 초과 문서는 오류가 아니라 앞부분만 사용해 정상 처리된다.
    빈 문서·추출 실패·권한 거부는 스킵 사유와 함께 반환되고, 호출자는 나머지 파일 처리를 계속한다.
    """
    try:
        text = extract_text(path, max_chars)
    except ExtractionError as exc:
        reason = (
            SkipReason.PERMISSION_DENIED
            if isinstance(exc.__cause__, PermissionError)
            else SkipReason.EXTRACTION_FAILED
        )
        return PreparedDocument(skipped=SkippedFile(path=path, reason=reason, detail=str(exc)))

    # 방어적 절단 — 추출기가 앞부분만 읽더라도 요약 입력은 반드시 max_chars 이하다.
    text = text[:max_chars]
    if not text.strip():
        return PreparedDocument(skipped=SkippedFile(path=path, reason=SkipReason.EMPTY_DOCUMENT))
    return PreparedDocument(text=text)


def extract_text(path: Path, max_chars: int) -> str:
    """`path`의 앞부분 텍스트를 최대 `max_chars`자까지 추출한다.

    Args:
        path: 입력 파일 경로.
        max_chars: 요약 입력으로 쓸 상한 (스펙 §4.2, 기본 12,000).

    Returns:
        추출된 평문. 내용이 없으면 빈 문자열 (호출자가 `empty_document`로 스킵).

    Raises:
        ExtractionError: 지원하지 않는 확장자이거나 읽기·파싱에 실패한 경우.
    """
    extractor = EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        # 스캐너가 이미 `SUPPORTED_EXTENSIONS`로 걸러 주므로 정상 경로에서는 닿지 않는다.
        # 스캐너를 우회한 직접 호출용 방어선이다 (스펙 §4.2).
        raise ExtractionError(f"지원하지 않는 확장자입니다: {path.suffix}")
    return extractor(path, max_chars)


def _open_failure_detail(path: Path) -> str:
    """열기에 실패한 오피스 파일의 스킵 detail 문구를 만든다 (스펙 §4.3.1).

    **원인 판정을 예외 메시지 문자열 매칭으로 하지 않는다** — 라이브러리 버전이 올라가면
    조용히 깨진다. 대신 파일 앞 8바이트라는 안정된 근거를 쓴다. 이 함수는 이미 열기에
    실패한 경로에서만 불리므로 정상 경로에는 비용이 없다.

    OLE 시그니처는 **암호화된 OOXML과 구형 이진 포맷(`.xls`/`.ppt`)을 구분하지 못한다.**
    근거가 두 원인을 가르지 못하므로 detail에 둘을 함께 적어 정직하게 둔다 — 어느 쪽이든
    사용자가 할 일은 같다(암호를 풀거나 최신 포맷으로 다시 저장).

    시그니처를 읽지 못하면(파일 소실·권한 거부) 기본 문구로 떨어진다. 여기서 새 예외를
    올리면 원래의 추출 실패가 다른 예외로 뒤덮여 스킵 사유 자체가 바뀐다.
    """
    label = path.suffix.lower().lstrip(".") or "파일"
    base = f"{label}를 열지 못했습니다: {path}"
    try:
        with path.open("rb") as stream:
            head = stream.read(len(OLE_SIGNATURE))
    except OSError:
        return base
    if head == OLE_SIGNATURE:
        return f"{base} — 암호화되었거나 구형 이진 포맷입니다"
    return base


def _extract_plaintext(path: Path, max_chars: int) -> str:
    """`.txt`/`.md`를 UTF-8로 읽되 디코딩 오류는 대체 문자로 흡수한다."""
    try:
        with path.open("r", encoding="utf-8", errors="replace") as stream:
            return stream.read(max_chars)
    except OSError as exc:
        raise ExtractionError(f"파일을 읽지 못했습니다: {path}") from exc


def _extract_docx(path: Path, max_chars: int) -> str:
    """`.docx` 문단을 순서대로 잇되 `max_chars`에 도달하면 누적을 멈춘다."""
    try:
        document = Document(str(path))
    except (PackageNotFoundError, OSError, ValueError, KeyError) as exc:
        raise ExtractionError(f"docx를 열지 못했습니다: {path}") from exc

    parts: list[str] = []
    accumulated = 0
    try:
        for paragraph in document.paragraphs:
            text = paragraph.text
            if not text:
                continue
            parts.append(text)
            accumulated += len(text) + 1
            if accumulated >= max_chars:
                break
    except (OSError, ValueError, KeyError) as exc:
        raise ExtractionError(f"docx 문단을 읽지 못했습니다: {path}") from exc

    return "\n".join(parts)[:max_chars]


def _extract_pdf(path: Path, max_chars: int) -> str:
    """`.pdf` 텍스트 레이어를 페이지 순서로 잇되 `max_chars`에 도달하면 멈춘다 (스펙 §4.1).

    이미지/OCR은 하지 않는다. 암호화 PDF(`reader.is_encrypted`)와 손상·파싱 실패는
    `ExtractionError`로 올려 호출자가 `extraction_failed`로 스킵하게 하고(암호화는 detail로
    사유를 남긴다), 텍스트 레이어가 없으면(스캔 이미지 PDF 등) 빈 문자열을 돌려주어 호출자의
    빈 문서 검사가 `empty_document`로 처리하게 한다 (스펙 §4.1·§5).
    """
    try:
        reader = PdfReader(str(path))
        encrypted = reader.is_encrypted
    except (OSError, PyPdfError, ValueError) as exc:
        raise ExtractionError(f"pdf를 열지 못했습니다: {path}") from exc

    if encrypted:
        # 복호화·OCR은 비목표(스펙 §2). 사유를 detail로 남겨 스킵 리포트에 노출한다.
        raise ExtractionError("암호화된 PDF")

    parts: list[str] = []
    accumulated = 0
    try:
        for page in reader.pages:
            text = page.extract_text() or ""
            if not text:
                continue
            parts.append(text)
            accumulated += len(text) + 1
            if accumulated >= max_chars:
                break
    except (PyPdfError, ValueError, KeyError) as exc:
        raise ExtractionError(f"pdf 텍스트를 읽지 못했습니다: {path}") from exc

    return "\n".join(parts)[:max_chars]


#: 엑셀·PPTX 개봉이 실패할 때 라이브러리가 올리는 예외들 (스펙 §4.3.1).
#: `PermissionError`는 `OSError`의 하위 타입이라 여기에 포함되며, `from exc`로 원인을 보존해
#: 호출자가 `permission_denied`와 `extraction_failed`를 계속 가를 수 있다.
_OFFICE_OPEN_ERRORS = (InvalidFileException, BadZipFile, OSError, ValueError, KeyError)


def _cell_text(value: object) -> str:
    """엑셀 셀 값 하나를 표기 문자열로 바꾼다 (스펙 §4.2 T2 표).

    `openpyxl`이 돌려주는 것은 문자열이 아니라 파이썬 객체다. `str()`에 그대로 맡기면
    `120000.0`·`2026-03-31 00:00:00` 같은 표기가 나오고, LLM이 요약문에 그대로 옮겨 적는다.

    엑셀 표시 형식(`number_format`)은 해석하지 않는다 — `read_only` 모드에서 제한적이고,
    통화·백분율까지 다루면 이번 슬라이스의 「텍스트 레이어만」 원칙을 넘어선다.
    """
    if value is None:
        # 빈 셀. 이 값이 「내용 없음」 판정의 기준이다 (§4.2 T1).
        return ""
    if isinstance(value, bool):
        # `bool`은 `int`의 하위 타입이므로 숫자 규칙보다 먼저 가른다.
        return str(value)
    if isinstance(value, datetime):
        if (value.hour, value.minute, value.second, value.microsecond) == (0, 0, 0, 0):
            return value.strftime("%Y-%m-%d")
        return value.strftime("%Y-%m-%d %H:%M")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _excel_row_line(row: Iterable[object]) -> str:
    """행 하나를 탭으로 이은 한 줄로 만든다. 전 셀이 비면 빈 문자열을 돌려준다 (§4.2 T2).

    행 **끝**의 빈 셀은 탭을 남기지 않는다(오른쪽 여백 탭 제거). 앞·중간의 빈 셀은 열 위치
    정보이므로 탭으로 남긴다.
    """
    cells = [_cell_text(value) for value in row]
    while cells and not cells[-1]:
        cells.pop()
    return "\t".join(cells)


def _extract_xlsx(path: Path, max_chars: int) -> str:
    """`.xlsx`/`.xlsm`의 셀 값을 시트 순서로 잇되 `max_chars`에 도달하면 멈춘다 (스펙 §4.2).

    `read_only=True`로 열어 통합문서를 통째로 메모리에 올리지 않는다. `data_only=True`는
    수식 대신 **캐시된 계산값**을 읽는다 — 요약에 쓸모 있는 것은 `=SUM(B2:B10)`이 아니라 그
    결과값이다. Excel이 저장하지 않은 파일에는 그 캐시가 없어 수식 셀이 비어 읽히며, 시트가
    수식뿐이면 추출 결과가 빈 문자열이 되어 호출자가 `empty_document`로 스킵한다 (§5).

    **경계 줄은 내용이 있는 시트에만 낸다** (§4.2 T1). 경계 줄은 「뒤에 내용이 온다」는
    신호이므로 뒤가 비면 낼 것이 없다. 항상 내면 값이 전부 비는 파일의 추출 결과가 경계
    줄만 남아 호출자의 `text.strip()` 검사를 통과해 버려, `empty_document`로 걸러져야 할
    문서가 시트 이름만 든 입력으로 LLM까지 간다. 빈 판정 책임은 호출자 한 곳에 그대로 둔다.

    **숨긴 시트는 제외하고 숨긴 행은 제외하지 않는다.** `sheet_state`는 읽기 전용 개봉에서도
    판정되지만, `ReadOnlyWorksheet`에는 `row_dimensions` 속성이 아예 없어 「이 행이 숨겨졌는가」를
    알 길이 없다 (2026-08-27 · `openpyxl` 3.1.5 실측 · 스펙 §4.2 T3). `read_only=False`로
    바꾸는 선택지는 위의 「통째로 메모리에 올리지 않는다」와 정면으로 부딪쳐 택하지 않는다.
    """
    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
    except _OFFICE_OPEN_ERRORS as exc:
        raise ExtractionError(_open_failure_detail(path)) from exc

    parts: list[str] = []
    accumulated = 0
    try:
        for worksheet in workbook.worksheets:
            if worksheet.sheet_state != "visible":
                continue
            # 첫 내용 줄을 만나기 전까지는 경계 줄을 내지 않고 들고만 있는다.
            pending_boundary: str | None = f"[시트: {worksheet.title}]"
            for row in worksheet.iter_rows(values_only=True):
                line = _excel_row_line(row)
                if not line:
                    continue
                if pending_boundary is not None:
                    parts.append(pending_boundary)
                    accumulated += len(pending_boundary) + 1
                    pending_boundary = None
                parts.append(line)
                accumulated += len(line) + 1
                if accumulated >= max_chars:
                    break
            if accumulated >= max_chars:
                break
    except (BadZipFile, OSError, ValueError, KeyError) as exc:
        raise ExtractionError(f"{path.suffix.lstrip('.')} 시트를 읽지 못했습니다: {path}") from exc
    finally:
        workbook.close()

    return "\n".join(parts)[:max_chars]


#: 확장자 → 추출기 디스패치 (스펙 §4.2). **`extract.py`의 지원 포맷 목록은 이 매핑 하나뿐이다** —
#: 확장자 상수를 따로 두지 않는다.
#:
#: `config.SUPPORTED_EXTENSIONS`와 **키 집합이 정확히 같아야 하며**, 그 정합성은
#: `tests/unit/test_extract.py`의 단위테스트가 단언한다. 두 정의를 한쪽에서 다른 쪽을 참조해
#: 파생시키지 않는 것은 의도적이다 — 파생시키면 단언이 공허해지고, 어긋남을 잡아낼 감시장치가
#: 사라진다. 정본은 `config.SUPPORTED_EXTENSIONS`이고 이 매핑은 「그 목록을 실제로 처리할 수
#: 있는가」를 독립적으로 진술한다.
#:
#: 값은 전부 `(path, max_chars) -> str` 시그니처를 공유한다. 포맷을 더할 때 손댈 곳은
#: 추출기 함수 하나와 이 매핑 한 줄, 그리고 `config.SUPPORTED_EXTENSIONS` 한 값이다.
EXTRACTORS: dict[str, Callable[[Path, int], str]] = {
    ".txt": _extract_plaintext,
    ".md": _extract_plaintext,
    ".docx": _extract_docx,
    ".pdf": _extract_pdf,
    ".xlsx": _extract_xlsx,
    ".xlsm": _extract_xlsx,
}
